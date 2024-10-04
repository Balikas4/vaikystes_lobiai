from django.shortcuts import render, redirect
from .forms import EmailForm
from django.contrib import messages
from about_us.models import AboutUsPage, Grupe, DailyRoutineActivity
from gallery.models import GalleryCategory
from main_page.models import MainReview, MainPage
from nutrition.models import NutritionPage, WeeklyNutrition
from admissions.models import LankymoKaina, NuolaidosIrKompensacijos
from register.models import Registration
from register.forms import RegistrationForm
from django.db.models import Case, When, Value, IntegerField
from contact.models import Contact
from django.core.mail import send_mail, BadHeaderError
from django.http import HttpResponse
from django.core.mail import EmailMessage
from django.core.files import File
from docx import Document
from django.utils import timezone
import os


def home(request):
    main_page = MainPage.objects.first()
    reviews = MainReview.objects.filter(main_page=main_page)
    return render(request, 'main.html', {'main_page': main_page, 'reviews': reviews})

def about(request):
    about_us_page = AboutUsPage.objects.first()  # Assuming there's only one AboutUsPage
    groups = Grupe.objects.filter(about_us_page=about_us_page)

    # Define days of the week in Lithuanian
    days_of_week = ["Pirmadienis", "Antradienis", "Trečiadienis", "Ketvirtadienis", "Penktadienis"]

    # Prepare routine data
    routine_data = {}
    for group in groups:
        group_routines = {day: None for day in days_of_week}
        for routine in group.routines.all():
            # Fetch activities ordered by the 'order' field
            activities = DailyRoutineActivity.objects.filter(dailyroutine=routine).order_by('order')
            # Prepare a list of activity objects in the correct order
            ordered_activities = [activity.activity for activity in activities]
            group_routines[routine.day] = {
                'routine': routine,
                'activities': ordered_activities
            }
        routine_data[group.id] = group_routines

    context = {
        'about_us_page': about_us_page,
        'groups': groups,
        'routine_data': routine_data,
        'days_of_week': days_of_week,
    }
    return render(request, 'about.html', context)

def education(request):
    return render(request, 'education.html')

def nutrition(request):
    # Fetch the NutritionPage
    nutrition_page = NutritionPage.objects.first()  # Adjust as needed to get the specific NutritionPage
    
    # Define the correct order of days
    days_order = ['Pirmadienis', 'Antradienis', 'Trečiadienis', 'Ketvirtadienis', 'Penktadienis']
    
    # Fetch Weekly Nutrition data, grouped by week and sorted by day order
    weeks = {}
    for week in range(1, 5):
        weeks[week] = WeeklyNutrition.objects.filter(
            week_number=week
        ).order_by(
            Case(
                *[When(day=day, then=Value(index)) for index, day in enumerate(days_order)],
                output_field=IntegerField()
            )
        )

    context = {
        'nutrition_page': nutrition_page,
        'weeks': weeks,
    }
    return render(request, 'nutrition.html', context)


def admissions(request):
    lankymo_kaina = LankymoKaina.objects.first()  # Adjust if there are multiple instances
    nuolaidos_ir_kompensacijos = NuolaidosIrKompensacijos.objects.first()  # Get the first instance

    context = {
        'lankymo_kaina': lankymo_kaina,
        'nuolaidos_ir_kompensacijos': nuolaidos_ir_kompensacijos
    }

    return render(request, 'admissions.html', context)


def gallery(request):
    categories = GalleryCategory.objects.all()
    context = {
        'categories': categories
    }
    return render(request, 'gallery.html', context)

def contact(request):
    context = {
        'contact': Contact.objects.first(),  # Example, adjust according to your context
    }
    return render(request, 'contact.html', context)

def replace_placeholders(doc, replacements):
    """Replace placeholders in the document with actual form data and optionally bold specific text."""

    for paragraph in doc.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        # Replace the text
                        run.text = run.text.replace(placeholder, value)


    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in replacements.items():
                    if placeholder in cell.text:
                        for run in cell.paragraphs[0].runs:
                            if placeholder in run.text:
                                # Replace the text
                                run.text = run.text.replace(placeholder, value)


def register(request):
    if request.method == 'POST':
        form = RegistrationForm(request.POST)
        if form.is_valid():
            registration = form.save()

            # Prepare admin email
            admin_subject = 'Naujas registracijos pateikimas'
            admin_message = f'''
            Naujas registracijos įrašas buvo pateiktas su šiais duomenimis:

            **Tėvų/Globėjų Informacija**
            - Vardas: {registration.first_last_name}
            - Kontaktinis Telefonas: {registration.contact_phone}
            - El. Paštas: {registration.email}
            - Namų Adresas: {registration.home_address}

            **Dokumento ir Priėmimo Datos**
            - Dokumento Data: {registration.document_date}
            - Vaiko Vardas: {registration.child_first_last_name}
            - Priėmimo Data: {registration.admission_date}

            **Vaiko Informacija**
            - Vaiko Vardas: {registration.child_first_name}
            - Vaiko Pavardė: {registration.child_last_name}
            - Vaiko Asmens Kodas: {registration.child_personal_code}
            - Vaiko Namų Adresas: {registration.child_home_address}

            **Tėvų Informacija**
            - Tėčio Informacija: {registration.father_info}
            - Mamos Informacija: {registration.mother_info}

            **Vaiko Sveikatos ir Gebėjimų Informacija**
            - Vaiko Sveikatos Informacija: {registration.child_health_info}
            - Vaiko Talentai: {registration.child_talents}
            '''

            # Load the existing document template
            template_path = '/app/register/word/Registracija.docx'
            doc = Document(template_path)

            # Define replacements based on form data
            replacements = {
                '{{first_last_name}}': registration.first_last_name,
                '{{contact_phone}}': registration.contact_phone,
                '{{email}}': registration.email,
                '{{home_address}}': registration.home_address,
                '{{document_date}}': str(registration.document_date),
                '{{child_first_last_name}}': registration.child_first_last_name,
                '{{admission_date}}': str(registration.admission_date),
                '{{child_first_name}}': registration.child_first_name,
                '{{child_last_name}}': registration.child_last_name,
                '{{child_personal_code}}': registration.child_personal_code,
                '{{child_home_address}}': registration.child_home_address,
                '{{father_info}}': registration.father_info,
                '{{mother_info}}': registration.mother_info,
                '{{child_health_info}}': registration.child_health_info,
                '{{child_talents}}': registration.child_talents,
            }

            # Replace placeholders with actual data
            replace_placeholders(doc, replacements)

            # Save the modified document to a temporary file
            doc_filename = f'Registracija_{timezone.now().strftime("%Y%m%d_%H%M%S")}.docx'
            doc_path = os.path.join('/tmp', doc_filename)
            doc.save(doc_path)

            # Save the document to the model
            with open(doc_path, 'rb') as doc_file:
                registration.document.save(doc_filename, File(doc_file), save=True)

            # Optionally delete the temporary file
            # Ensure this file is removed after it's used to avoid clutter
            try:
                # Prepare admin email with attachment
                email = EmailMessage(
                    subject=admin_subject,
                    body=admin_message,
                    from_email='vaikysteslobiaiweb@gmail.com',
                    to=['rimsa.mindaugas@gmail.com'],
                )

                # Attach the document
                email.attach_file(doc_path)
                email.send()

                # Prepare welcome email
                welcome_email = EmailMessage(
                    subject='Jūsų prašymas buvo gautas - Vaikystės lobiai',
                    body=(
                        'Dėkojame,\n\n'
                        'Jūsų prašymas buvo gautas, netrukus susisieksime.\n\n'
                        'Pagarbiai,\n'
                        '"vaikystės lobiai" administracija'
                    ),
                    from_email='vaikysteslobiaiweb@gmail.com',
                    to=[registration.email],
                )
                welcome_email.send()

                # Optionally delete the temporary file
                os.remove(doc_path)

                # Success message and redirect
                messages.success(request, 'Sėkmingai užpildėte registracijos prašymą.')
                return redirect('register')

            except BadHeaderError:
                return HttpResponse('Invalid header found.')
            except Exception as e:
                # Log the error message for further inspection
                print(f"Error sending email: {e}")
                messages.error(request, 'Įvyko klaida siunčiant el. laišką. Bandykite dar kartą.')

        else:
            messages.error(request, 'Įvyko klaida registruojantis.')
    else:
        form = RegistrationForm()

    return render(request, 'register.html', {'form': form})
